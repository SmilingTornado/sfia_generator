# Create your views here.
import docx
import gensim
import numpy as np
from django.conf import settings
from django.http import HttpResponse
from django.shortcuts import render
from docx.shared import RGBColor, Inches, Pt
from nltk.tokenize import sent_tokenize, word_tokenize

from .models import Skill, Level


# View for home page
def index(request):
    # Request to get the form
    if request.method == "GET":
        context = {'searched': False}
        # Return the standard blank form
        return render(request, 'form.html', context)
    # POSTing to the form means the request body will have some data
    elif request.method == "POST":
        # Checks whether the POST request wants to generate a form by checking request body
        if 'type' in request.POST and 'sk1' in request.POST and 'sk2' in request.POST \
                and 'sk1_min' in request.POST and 'sk2_min' in request.POST \
                and 'sk1_max' in request.POST and 'sk2_max' in request.POST:
            # Checking validity of request
            if is_valid(request):
                # Generate and return the form
                return generate(request)
            else:
                # Return the page for an invalid request
                return render(request, 'invalid.html', {})

        # If data was posted from the search function form.
        elif 'input' in request.POST:
            return search_similarities(request)

        # If data was posted from the skill selector
        elif 'code_1' and 'code_2' in request.POST:
            context = {'searched': False, 'sk1_code': request.POST['code_1'], 'sk2_code': request.POST['code_2']}
            return render(request, 'form.html', context)

        # Any other type of POST request would be invalid
        else:
            return render(request, 'invalid.html', {})
    else:
        # Any other request would just be returned the blank form
        context = {'searched': False}
        return render(request, 'form.html', context)


# View for search page
def search_page(request):
    # Returns the search page
    return render(request, 'search.html', {})


# View to list skills
def list_skills(request):
    set_1, set_2, set_3 = get_skill_sets()  # Gets skills in 3 evenly split sets
    return render(request, 'list_skills.html', {"set_1": set_1, "set_2": set_2,
                                                "set_3": set_3})  # Renders and returns the page of the list of skills


# View to list skills for second skill selection
def select_second(request, code_1):  # Same as list_skills but addional context is added to be rendered
    set_1, set_2, set_3 = get_skill_sets()  # Gets skills in 3 evenly split sets
    return render(request, 'list_skills.html', {"code_1": code_1, "set_1": set_1, "set_2": set_2,
                                                "set_3": set_3})  # Renders and returns the page of the list of skills


def get_skill_sets():
    set_1 = []  # Column 1
    set_2 = []  # Column 2
    set_3 = []  # Column 3
    skill_objects = Skill.objects.all().order_by('code')  # Get all the skills and order them by the skill code
    length = len(skill_objects)  # Find number of skills
    for num, skill in enumerate(skill_objects, start=0):
        if num < length / 3:  # Checks if the skill is in the first third of the list
            set_1.append(skill)  # Appends to first column set
        elif num < length * (2 / 3):  # Checks if the skill is in the second third of the list
            set_2.append(skill)  # Appends to the second column set
        else:  # All other skills
            set_3.append(skill)  # Appended to the last column set
    return set_1, set_2, set_3


# View details of skill
def show_skill(request, code):
    try:
        skill_object = Skill.objects.get(code=code.lower())  # Get the skill from the code
        levels = Level.objects.filter(skill=skill_object)  # Get the levels using the skill_object as the key
        context = {
            'skill': skill_object,
            'levels': levels
        }  # Prepare context for rendering onto template
        return render(request, 'show_skill.html', context)  # Render and return context
    except:  # In the case where the skill code is invalid
        return render(request, 'invalid.html', {})  # Return page for invalid requests


#View details of second selected skill
def view_second(request, code_1, code_2):
    try:
        skill_object = Skill.objects.get(code=code_2.lower())  # Get the skill from the code
        levels = Level.objects.filter(skill=skill_object)  # Get the levels using the skill_object as the key
        context = {
            'skill': skill_object,
            'levels': levels,
            'code_1': code_1,
            'code_2': code_2
        }  # Prepare context for rendering onto template
        return render(request, 'show_skill.html', context)  # Render and return context
    except:  # In the case where the skill code is invalid
        return render(request, 'invalid.html', {})  # Return page for invalid requests


def search_similarities(request):
    similarities = {}  # Dictionary to store the calculated similarities
    input = request.POST['input']  # Get the input from the request
    # Create a list of sentences where each sentence has been broken down into a list of words
    gen_docs = [[w.lower() for w in word_tokenize(text)]
                for text in sent_tokenize(input)]
    # Create a dictionary of unique words
    dictionary = gensim.corpora.Dictionary(gen_docs)
    # Generate bag of words to measure frequency of word use
    corpus = [dictionary.doc2bow(gen_doc) for gen_doc in gen_docs]
    # Calculate Term Frequency, Inverse Document Frequency of words
    tf_idf = gensim.models.TfidfModel(corpus)
    # Create similarity model
    sims = gensim.similarities.Similarity(settings.BASE_DIR + '/Generator/gensim', tf_idf[corpus],
                                          num_features=len(dictionary))
    # Checking for similarities with level descriptions
    for level in Level.objects.all():
        skill_sim_total = 0

        for sentence in sent_tokenize(level.description):
            query_doc = [w.lower() for w in word_tokenize(sentence)]
            query_doc_bow = dictionary.doc2bow(query_doc)
            query_doc_tf_idf = tf_idf[query_doc_bow]
            sum_of_sims = (np.sum(sims[query_doc_tf_idf], dtype=np.float32))
            similarity = float(sum_of_sims / len(sent_tokenize(input)))
            skill_sim_total += similarity

        skill_sim = skill_sim_total / len(sent_tokenize(level.description))
        # Check if similarities for a skill has been calculated before
        if level.skill.code not in similarities:
            similarities[level.skill.code] = skill_sim
        # If calculated before, check if new description is more similar
        elif similarities[level.skill.code] < skill_sim:
            similarities[level.skill.code] = skill_sim
    # Checking for similarities with skill descriptions
    # Same procedure as with for levels
    for skill in Skill.objects.all():
        skill_sim_total = 0

        for sentence in sent_tokenize(skill.description):
            query_doc = [w.lower() for w in word_tokenize(sentence)]
            query_doc_bow = dictionary.doc2bow(query_doc)
            query_doc_tf_idf = tf_idf[query_doc_bow]
            sum_of_sims = (np.sum(sims[query_doc_tf_idf], dtype=np.float32))
            similarity = float(sum_of_sims / len(sent_tokenize(input)))
            skill_sim_total += similarity

        skill_sim = skill_sim_total / len(sent_tokenize(skill.description))
        if skill.code not in similarities:
            similarities[skill.code] = skill_sim
        elif similarities[skill.code] < skill_sim:
            similarities[skill.code] = skill_sim
    # Find the most similar skill
    first_match = max(similarities, key=similarities.get)
    # If the maximum similarity score was 0, return the form
    if (similarities[first_match] == 0):
        return render(request, 'form.html', {'searched': True})
    # Removes the most similar skill
    similarities.pop(first_match, None)
    # Finds the current maximum similarity score
    second_match = max(similarities, key=similarities.get)
    # If the new maximum similarity score is 0, return only the first match
    if (similarities[second_match] == 0):
        return render(request, 'form.html', {'sk1_code': first_match.upper, 'searched': True})
    # Return rendered form with found matches
    context = {'sk1_code': first_match.upper, 'sk2_code': second_match.upper, 'searched': True}
    return render(request, 'form.html', context)


# Returns whether a skill is valid
def is_valid(request):
    # Grabbing data from request
    sk1 = request.POST['sk1']
    sk1_start = int(request.POST['sk1_min'])
    sk1_stop = int(request.POST['sk1_max'])
    sk2 = request.POST['sk2']
    sk2_start = int(request.POST['sk2_min'])
    sk2_stop = int(request.POST['sk2_max'])
    type = request.POST['type']
    # Check if request is valid
    if 'type' in request.POST and 'sk1' in request.POST and 'sk2' in request.POST \
            and 'sk1_min' in request.POST and 'sk2_min' in request.POST \
            and 'sk1_max' in request.POST and 'sk2_max' in request.POST:
        if sk1_start >= 1 and sk2_start >= 1 and sk1_stop <= 7 and sk2_stop <= 7 and (
                type == 'student' or type == 'employer'):
            try:  # Try to retrieve the skill object
                skill_object = Skill.objects.get(code=sk1.lower())
            except:
                return False
            if sk2 != '':  # If the second skill isn't blank
                try:  # Try to retrieve the second skill object
                    skill_object = Skill.objects.get(code=sk2.lower())
                except:
                    return False
            return True
        else:
            return False
    else:
        return False


def generate(request):
    # Setting variables taken from request body
    sk1 = request.POST['sk1']
    sk1_start = int(request.POST['sk1_min'])
    sk1_stop = int(request.POST['sk1_max'])
    sk2 = request.POST['sk2']
    sk2_start = int(request.POST['sk2_min'])
    sk2_stop = int(request.POST['sk2_max'])
    type = request.POST['type']
    dedicate = False
    # Check if skills are to be rendered on dedicated pages
    if 'dedicate' in request.POST:
        dedicate = True

    # Generating the document
    if type == 'employer':
        doc = docx.Document(settings.BASE_DIR + '/Generator/DocxTemplates/employer_template.docx')
    else:
        doc = docx.Document(settings.BASE_DIR + '/Generator/DocxTemplates/student_template.docx')
    if dedicate:
        # Addidng a page break
        add_page_break(doc)
    if sk2 != '':
        sk1_concat = ''.join([level['description'] for level in get_levels(sk1, [sk1_start, sk1_stop])])
        sk2_concat = ''.join([level['description'] for level in get_levels(sk2, [sk2_start, sk2_stop])])
        # Check if skill 1 is longer than skill 2
        if len(sk1_concat) <= len(sk2_concat):
            # Adding skill information
            add_skill_info(sk1, doc)
            # Adding the first table
            add_skill_table(sk1, [sk1_start, sk1_stop], doc)
            # Addidng a page break
            add_page_break(doc)
            # Adding skill information
            add_skill_info(sk2, doc)
            # Adding the second table
            add_skill_table(sk2, [sk2_start, sk2_stop], doc)
            filename = '%s-%s.docx' % (sk1.upper(), sk2.upper())
        else:
            # Adding skill information
            add_skill_info(sk2, doc)
            # Adding the first table
            add_skill_table(sk2, [sk2_start, sk2_stop], doc)
            # Addidng a page break
            add_page_break(doc)
            # Adding skill information
            add_skill_info(sk1, doc)
            # Adding the second table
            add_skill_table(sk1, [sk1_start, sk1_stop], doc)
            filename = '%s-%s.docx' % (sk2.upper(), sk1.upper())
    else:
        # Adding skill information
        add_skill_info(sk1, doc)
        # Adding the first table
        add_skill_table(sk1, [sk1_start, sk1_stop], doc)
        filename = '%s.docx' % (sk1.upper())

    # Saving to output
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=' + filename
    doc.save(response)
    return response


# Get skill information
def get_skill(sk_code):
    skill_object = Skill.objects.get(code=sk_code.lower())
    # Put skill information into dictionary
    skill = {
        'name': skill_object.name,
        'code': skill_object.code,
        'description': skill_object.description,
        'levels': []
    }
    # Put each level's information into a dictionary and append to levels list in the skills dictionary
    for level in Level.objects.filter(skill=skill_object):
        skill['levels'].append({
            'level': level.level,
            'description': level.description,
        })
    # Return the dictionary
    return skill


# Get levels in a certain range
def get_levels(sk_code, sk_range):
    sk = get_skill(sk_code)
    levels = []
    # Put each level's information into a dictionary and append to levels list in the skills dictionary
    for i in range(sk_range[0], sk_range[1] + 1):
        for level in sk['levels']:
            if level['level'] == i:
                description = level['description']
                levels.append({'level': i, 'description': description})
                break
    return levels


def add_skill_table(sk_code, sk_range, doc):
    # Get the information for the skill
    levels = get_levels(sk_code, sk_range)

    # Table Generation
    t = doc.add_table(2, len(levels))  # Create Table
    t.autofit = True
    t.style = 'Table Grid'
    t.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER

    # Finding total length of descriptions for width calculations later
    total_description_length = 0
    for level in levels:
        total_description_length += len(level["description"])

    # Populating cells
    cell_count = 0
    for level in levels:
        top_cell = t.cell(0, cell_count).paragraphs[0].add_run('Level ' + str(level['level']))
        top_cell.bold = True
        top_cell.font.name = 'Calibri'
        bottom_cell = t.cell(1, cell_count).paragraphs[0].add_run(level['description'])
        bottom_cell.font.name = 'Calibri'
        bottom_cell.font.size = Pt(10)
        cell_width = 1.25 / len(levels) + 10.75 * len(level['description']) / total_description_length
        t.cell(0, cell_count).width = Inches(cell_width)
        t.cell(1, cell_count).width = Inches(cell_width)
        cell_count += 1


# Generate description for the skill
def add_skill_info(sk_code, doc):
    sk = get_skill(sk_code)
    p = doc.add_paragraph('')
    name = p.add_run(sk['name'] + ' ')
    name.bold = True
    name.font.size = Pt(14)
    name.font.name = 'Calibri'
    code = p.add_run(sk['code'].upper())
    code.bold = True
    code.font.size = Pt(11)
    code.font.color.rgb = RGBColor(0x89, 0x89, 0x89)
    code.font.name = 'Calibri'
    description = p.add_run(' – ' + sk['description'])
    description.font.size = Pt(10)
    description.font.name = 'Calibri'


def add_page_break(doc):
    paragraph = doc.add_paragraph('')
    run = paragraph.add_run('')
    run.add_break(docx.enum.text.WD_BREAK.PAGE)
